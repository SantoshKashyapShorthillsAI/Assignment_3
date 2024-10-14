from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Create a new Document
doc = Document()

# Add a title
doc.add_heading('Document with Links, Text, Table, and Images', 0)

# Add some text
doc.add_paragraph("This is a Python-generated Word document with multiple elements.")

# Function to add a hyperlink (fix applied)
def add_hyperlink(paragraph, url, text):
    # Create the hyperlink element
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    # Create the hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    # Create a run with the hyperlink text
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Styling for hyperlink (optional)
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    
    run.append(rPr)
    run.text = text
    
    # Add the hyperlink run to the hyperlink element
    hyperlink.append(run)
    
    # Append hyperlink to the paragraph
    paragraph._p.append(hyperlink)

# Adding a paragraph with a hyperlink
p = doc.add_paragraph("For more information, visit: ")
add_hyperlink(p, "https://www.python.org", "Python's official website")

# Add a table
doc.add_heading('Table Example', level=1)

# Create a table with 2 rows and 2 columns
table = doc.add_table(rows=2, cols=2)

# Add column headers
table.cell(0, 0).text = 'Name'
table.cell(0, 1).text = 'Age'

# Add row data
table.cell(1, 0).text = 'John'
table.cell(1, 1).text = '30'

# Add an image (Make sure the image is in the same directory as the script)
doc.add_heading('Image Example', level=1)
image_path = 'apple.jpeg' # Update with your image path
doc.add_picture(image_path, width=Inches(4))

# Save the document
doc.save('docx_example.docx')
print("Document created successfully!")
