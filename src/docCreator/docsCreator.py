from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Create a new Document
doc = Document()

# Add a title
doc.add_heading('Document with Main Text and Comments', 0)

# Add some main text
p = doc.add_paragraph("This is the main content of the document. ")
p.add_run("Some parts of the content might require comments. ")

# Function to simulate comments (in-line using a different style)
def add_comment(paragraph, comment_text):
    run = paragraph.add_run(f"[Comment: {comment_text}]")
    run.italic = True  # Italicize the comment
    run.font.size = Pt(10)  # Set font size smaller for comments
    run.font.color.rgb = RGBColor(255, 0, 0)  # Set comment color to red

# Adding comments to the document
p = doc.add_paragraph("Here's a section that might need a comment. ")
add_comment(p, "This is an example of a comment added in the document.")

# Adding another comment
p = doc.add_paragraph("Another piece of text with a comment at the end.")
add_comment(p, "This comment provides additional clarification.")

# Save the document
doc.save('/home/shtlp_0103/Assignment_3/Documents/doc_with_comments.docx')

print("Document with simulated comments created successfully!")
