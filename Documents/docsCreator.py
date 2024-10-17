from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Create a new Document
doc = Document()

# Add a title
doc.add_heading('Document with Links, Text, Complex Tables, and Footnotes', 0)

# Add some introductory text
doc.add_paragraph("This is a Python-generated Word document with multiple elements including complex tables and footnotes.")

# Function to add a hyperlink
def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    run.append(rPr)
    run.text = text
    hyperlink.append(run)
    paragraph._p.append(hyperlink)

# Adding a paragraph with a hyperlink
p = doc.add_paragraph("For more information, visit: ")
add_hyperlink(p, "https://www.python.org", "Python's official website")

# Add a footnote
p.add_run(" [1]")  # Reference in the text
footnote = doc.add_paragraph()
footnote.add_run("1. Python is a programming language that lets you work quickly and integrate systems more effectively.")

# Add a section heading for complex table
doc.add_heading('Complex Table Example', level=1)

# Create a complex table with merged cells
table = doc.add_table(rows=4, cols=3)

# Merge cells for the table title
cell = table.cell(0, 0)
cell.merge(table.cell(0, 2))  # Merge first row across three columns
cell.text = "Merged Header (spanning 3 columns)"
cell.paragraphs[0].runs[0].font.bold = True  # Bold text

# Fill in table headers
table.cell(1, 0).text = 'Column 1'
table.cell(1, 1).text = 'Column 2'
table.cell(1, 2).text = 'Column 3'

# Add data to the table
table.cell(2, 0).text = 'Row 1 Col 1'
table.cell(2, 1).text = 'Row 1 Col 2'
table.cell(2, 2).text = 'Row 1 Col 3'

table.cell(3, 0).text = 'Row 2 Col 1'
table.cell(3, 1).text = 'Row 2 Col 2'
table.cell(3, 2).text = 'Row 2 Col 3'

# Add an image (Ensure the image path is correct)
doc.add_heading('Image Example', level=1)
image_path = '/home/shtlp_0103/Assignment_3/Documents/apple.jpeg'  # Update with your image path
doc.add_picture(image_path, width=Inches(4))

# Save the document
doc.save('/home/shtlp_0103/Assignment_3/Documents/footnotes.docx')
print("Document with footnotes created successfully!")
